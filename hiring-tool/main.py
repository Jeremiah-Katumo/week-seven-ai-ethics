import streamlit as st
import pandas as pd
import numpy as np
from sklearn.model_selection import train_test_split
from sklearn.linear_model import LogisticRegression
from sklearn.tree import DecisionTreeClassifier
from sklearn.ensemble import RandomForestClassifier
from sklearn.preprocessing import LabelEncoder
from fairlearn.reductions import ExponentiatedGradient, DemographicParity, EqualizedOdds
from fairlearn.metrics import MetricFrame, selection_rate, demographic_parity_difference, equalized_odds_difference
from sklearn.metrics import accuracy_score, confusion_matrix, classification_report
import matplotlib.pyplot as plt
import io
import pdfplumber
import docx
from reportlab.pdfgen import canvas
import sqlalchemy
from sqlalchemy import create_engine, text
import spacy
import re
# import add_logo


# Set up the Streamlit app configuration
st.set_page_config(page_title="Fair AI Hiring Tool", page_icon=":briefcase:", layout="wide")
st.title("🤖 Fair AI Hiring Tool")


# Load spaCy model
@st.cache_resource
def load_nlp():
    return spacy.load("en_core_web_sm")

nlp = load_nlp()

engine = create_engine('mysql+pymysql://root:Mysql.003@localhost/hiring_tool')  
connection = engine.connect()

connection.execute(text("""
CREATE TABLE IF NOT EXISTS reports (
    id INT AUTO_INCREMENT PRIMARY KEY,
    accuracy FLOAT NOT NULL,
    metric_constraint VARCHAR(255) NOT NULL,
    metric_name VARCHAR(255) NOT NULL,
    metric_value FLOAT NOT NULL,
    created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP
)
"""))
connection.commit()


# --- Helper: Extract features from job/resume text ---
def extract_features_from_text(text):
    skills_list = ['python', 'java', 'sql', 'data', 'machine learning', 'deep learning',
                   'cloud', 'aws', 'azure', 'docker', 'kubernetes', 'pandas', 'numpy']
    degree_keywords = ['bachelor', 'master', 'phd', 'mba']
    experience_years = re.findall(r'(\d+)\+?\s*(years|yrs)', text.lower())

    skills_count = sum(1 for skill in skills_list if skill in text.lower())
    degree_score = max((i + 1 for i, d in enumerate(degree_keywords) if d in text.lower()), default=0)
    years_experience = max((int(y[0]) for y in experience_years), default=0)

    return [skills_count, degree_score, years_experience]


st.sidebar.header("📁 Upload Hiring Dataset")
uploaded_file = st.sidebar.file_uploader("Upload CSV, PDF, or DOCX file",
                                         type=["csv", "pdf", "docx"])

st.sidebar.header("⚙️ Model Configuration")
model_type = st.sidebar.selectbox("Select Model Type", ["Logistic Regression", "Decision Tree", "Random Forest"])
fairness_constraint = st.sidebar.selectbox("Select Fairness Constraint",
                                           ["Demographic Parity", "Equalized Opportunity"])

st.sidebar.header("📊 Model Evaluation")
evaluate_model = st.sidebar.button("Evaluate Model")
# if uploaded_file is not None:
#     if uploaded_file.type == "text/csv":
#         data = pd.read_csv(uploaded_file)
#     elif uploaded_file.type == "application/pdf":
#         with pdfplumber.open(uploaded_file) as pdf:
#             data = pd.DataFrame([page.extract_text() for page in pdf.pages])
#     elif uploaded_file.type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
#         doc = docx.Document(uploaded_file)
#         data = pd.DataFrame([para.text for para in doc.paragraphs])
#     else:
#         st.error("Unsupported file type. Please upload a CSV, PDF, or DOCX file.")
        
# Resume Upload Section
with st.expander("Upload and Parse Resumes (PDF/DOCS)", expanded=True):
    resume_file = st.file_uploader("Upload Resume", type=["pdf", "docx"])
    if resume_file is not None:
        resume_text = ""
        if resume_file.type == "application/pdf":
            with pdfplumber.open(resume_file) as pdf:
                resume_text = "\\n".join([page.extract_text() for page in pdf.pages])
        elif resume_file.type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
            doc = docx.Document(resume_file)
            resume_text = "\\n".join([para.text for para in doc.paragraphs])
        st.text_area("Resume Text", value=resume_text, height=200)
        st.info("Note: Resume Scoring is based on mock data features. Add NLP model for real scoring.")
        

if uploaded_file is not None and evaluate_model:
    data = pd.read_csv(uploaded_file) if uploaded_file.type == "text/csv" else pd.DataFrame()
    
    st.subheader("📊 Dataset Overview")
    if all(col in data.columns for col in [
        "Title", "location", "Posting_date", "DESCRIPTION", "BASIC QUALIFICATIONS", "PREFERRED QUALIFICATIONS"
    ]):
        # Combine relevant text columns for feature extraction
        data['full_text'] = (
            data['Title'].astype(str) + " " +
            data['DESCRIPTION'].astype(str) + " " +
            data['BASIC QUALIFICATIONS'].astype(str) + " " +
            data['PREFERRED QUALIFICATIONS'].astype(str)
        )

        # Extract features for all rows
        feature_matrix = data['full_text'].apply(extract_features_from_text).tolist()
        X = np.array(feature_matrix)
        y = data['Title'].values  # Target is now Title

        # Encode job titles
        title_encoder = LabelEncoder()
        y_encoded = title_encoder.fit_transform(y)

        # Train/test split
        X_train, X_test, y_train, y_test = train_test_split(X, y_encoded, test_size=0.2, random_state=42)
        title_model = LogisticRegression(solver='liblinear', multi_class='ovr', max_iter=1000)
        title_model.fit(X_train, y_train)
    else:
        title_model = None
        title_encoder = None
        st.warning("Dataset does not contain all required columns for job title matching.")
    
    st.subheader("📈 Preview Dataset")
    st.dataframe(data.head(10))
    
    protected_feature_attrs = [col for col in data.columns if data[col].nunique() <= 10 and data[col].dtype == 'object']
    selected_feature_attrs = st.multiselect("🔐 Select Protected Feature Attributes", protected_feature_attrs)
    
    try:
        for attr in selected_feature_attrs:
            if attr in data.columns:
                st.subheader(f"📊 {attr} Distribution")
                st.bar_chart(data[attr].value_counts())
            
            st.markdown(f"Fairness Analysis for {attr}")
            
            X = data.drop(columns=[attr, 'Title']) # Assuming 'hired' is the label column or target column
            y = data[attr]
            
            X_train, X_test, y_train, y_test, s_train, s_test = train_test_split(X, y, test_size=0.3, random_state=42)
            
            base_model = None
            if model_type == "Logistic Regression":
                base_model = LogisticRegression(solver='liblinear')
            elif model_type == "Decision Tree":
                base_model = DecisionTreeClassifier(criterion='entropy', splitter='best', max_depth=10, random_state=42)
            elif model_type == "Random Forest":
                base_model = RandomForestClassifier(n_estimators=100, max_depth=10, random_state=42, criterion='entropy')
            else:
                st.error("Unsupported model type selected.")
                base_model = None
                
            base_model.fit(X_train, y_train)
                
            constraint = None
            fairness_value = None
            metric_label = None
            if fairness_constraint == "Demographic Parity":
                constraint = DemographicParity()
                fairness_value = demographic_parity_difference(
                    y_true=y_test, 
                    y_pred=base_model.predict(X_test), 
                    sensitive_features=s_test)
                metric_label = "Demographic Parity Difference"
            elif fairness_constraint == "Equalized Opportunity":
                constraint = EqualizedOdds()
                fairness_value = equalized_odds_difference(
                    y_true=y_test, 
                    y_pred=base_model.predict(X_test), 
                    sensitive_features=s_test)
                metric_label = "Equalized Odds Difference"
            elif fairness_constraint == "None":
                st.warning("No fairness constraint selected. Model will not be mitigated for fairness.")
                fairness_value = None
                metric_label = "None"
            else:
                st.error("Unsupported fairness constraint selected.")
                constraint = None
                fairness_value = None
                metric_label = "None"
                
            mitigator = ExponentiatedGradient(base_model, constraints=constraint, sensitive_features=s_train)
            mitigator.fit(X_train, y_train)
            
            y_preds = mitigator.predict(X_test)
            accuracy = accuracy_score(y_test, y_preds)
            conf_matrix = confusion_matrix(y_test, y_preds)
            class_report = classification_report(y_test, y_preds)
            metric_frame = MetricFrame(
                metrics={
                    "selection_rate": selection_rate,
                    "demographic_parity_difference": demographic_parity_difference,
                    "equalized_odds_difference": equalized_odds_difference
                },
                y_true=y_test,
                y_pred=y_preds,
                sensitive_features=s_test
            )
            
            if metric_frame.by_group.empty:
                st.warning("No data available for the selected protected feature attributes.")
            else:
                st.subheader("Fairness Metrics")
                st.dataframe(metric_frame.by_group)
                
                st.subheader("Model Accuracy")
                st.write(f"Accuracy: {accuracy:.2f}")
                
                st.subheader("Fairness Value")
                if fairness_value is not None:
                    st.write(f"{metric_label}: {fairness_value:.2f}")
                
                st.subheader("Confusion Matrix")
                st.write(conf_matrix)
                
                st.subheader("Classification Report")
                st.text(class_report)
                
                # Save report to database
                for metric_name, metric_value in metric_frame.by_group.items():
                    connection.execute("""
                    INSERT INTO reports (accuracy, metric_constraint, metric_name, metric_value) 
                    VALUES (%s, %s, %s, %s)
                    """, (accuracy, fairness_constraint, metric_name, metric_value))
                connection.commit()
                
                col_one, col_two = st.columns(2)
                with col_one:
                    st.subheader("Confusion Matrix Heatmap")
                    fig, ax = plt.subplots()
                    ax.matshow(conf_matrix, cmap='Blues', alpha=0.7)
                    for (i, j), z in np.ndenumerate(conf_matrix):
                        ax.text(j, i, z, ha='center', va='center')
                    plt.xlabel('Predicted')
                    plt.ylabel('True')
                    ax.set_title('Confusion Matrix')
                    ax.set_xticklabels([''] + list(set(y_test)))
                    ax.set_yticklabels([''] + list(set(y_test)))
                    st.pyplot(fig)
                    
                fig, ax = plt.subplots()
                metric_frame.by_group.plot.bar(ax=ax, color=['#00BFC4', '#F8766D'])
                ax.set_title('Fairness Metrics by Group')
                ax.set_ylabel('Selection Rate')
                ax.set_xlabel(f'Selection Rate by {attr}')
                st.subheader("Fairness Metrics Bar Chart")
                st.pyplot(fig)
            
            # Export report to PDF
            st.subheader("Export Fairness Analysis Report to PDF")
            if st.button("Export/Generate PDF Report"):
                pdf_buffer = io.BytesIO()
                pdf_canvas = canvas.Canvas(pdf_buffer)
                y_pos = 800  # or another suitable value
                pdf_canvas.drawString(100, 800, "Fairness Analysis Report")
                y_pos -= 30
                for attr in selected_feature_attrs:
                    pdf_canvas.drawString(100, y_pos, f"Protected Feature Attributes: {attr}")
                    y_pos -= 20
                pdf_canvas.drawString(100, y_pos, f"Model Type: {model_type}")
                y_pos -= 20
                pdf_canvas.drawString(100, y_pos, f"Fairness Constraint: {fairness_constraint}")
                y_pos -= 20
                pdf_canvas.drawString(100, y_pos, f"Accuracy: {accuracy:.2f}")
                y_pos -= 20
                pdf_canvas.drawString(100, y_pos, f"{metric_label}: {fairness_value:.2f}" if fairness_value is not None else "No fairness value calculated")
                
                # Add confusion matrix
                pdf_canvas.drawString(100, 650, "Confusion Matrix:")
                for i in range(len(conf_matrix)):
                    for j in range(len(conf_matrix[i])):
                        pdf_canvas.drawString(120 + j * 50, 630 - i * 20, str(conf_matrix[i][j]))
                
                # Add classification report
                pdf_canvas.drawString(100, 580, "Classification Report:")
                lines = class_report.split('\n')
                for i, line in enumerate(lines):
                    pdf_canvas.drawString(100, 560 - i * 20, line)
                
                pdf_canvas.save()
                
                pdf_buffer.seek(0)
                st.download_button("Download PDF Report", data=pdf_buffer, file_name="fairness_analysis_report.pdf", mime="application/pdf")
                
            # Real Time Resume Scoring Section
            st.subheader("Real Time Resume Job Title Matching (NLP-Based)")
            with st.form("resume_title_matching_form"):
                submit_button = st.form_submit_button("Match Resume to Job Titles")
                if submit_button:
                    if 'resume_text' in locals() and resume_text.strip() and title_model is not None:
                        features = extract_features_from_text(resume_text)
                        probs = title_model.predict_proba([features])[0]
                        top_n = 5
                        top_indices = np.argsort(probs)[::-1][:top_n]
                        top_titles = title_encoder.inverse_transform(top_indices)
                        top_scores = probs[top_indices] * 100  # Convert to percentage

                        st.write("### 🏆 Top Job Title Matches")
                        for title, score in zip(top_titles, top_scores):
                            st.write(f"**{title}**: {score:.2f}% match")
                    else:
                        st.warning("Please upload and parse a resume and ensure the model is trained.")
                        
    except Exception as e:
        st.error(f"Error processing the uploaded file: {e}")
        
    with st.expander("View Database Reports", expanded=True):
        st.subheader("Reports from Database")
        query = "SELECT * FROM reports ORDER BY created_at DESC"
        reports = pd.read_sql(query, connection)
        st.dataframe(reports)
        
else:
    st.warning("Please upload a dataset to proceed with model evaluation.")
                
    
    
    
# # Save to final app file
# with open("/mnt/data/fair_hiring_final_app.py", "w") as f:
#     f.write(final_app_code)
